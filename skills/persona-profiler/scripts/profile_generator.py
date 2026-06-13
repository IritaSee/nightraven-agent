#!/usr/bin/env python3
import os
import json
import argparse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set explicit margins (padding) for a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def apply_text_formatting(run, font_name="Times New Roman", size_pt=11, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def clean_text(text):
    """Replaces em-dashes with commas or parentheses as per user preference."""
    if not text:
        return ""
    # Replace em-dashes with commas or parentheses
    text = text.replace("—", ", ")
    text = text.replace("--", ", ")
    return text

def create_docx_profile(data, output_path):
    doc = Document()
    
    # Set Margins to 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set default style to Times New Roman 11pt, 1.15 line spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f"DEEP BEHAVIORAL PROFILE & PERSONA ENGINE\nCharacter: {data.get('character_name', 'Unnamed Character')}")
    apply_text_formatting(title_run, size_pt=14, bold=True)
    title_p.paragraph_format.space_after = Pt(18)
    
    # Introduction / Metadata
    intro_p = doc.add_paragraph()
    intro_run = intro_p.add_run("This document contains a high-fidelity behavioral profile and an LLM-compatible system prompt designed for realistic fictional character simulation. All personal identifiable information (PII) has been omitted or generalized to protect privacy.")
    apply_text_formatting(intro_run, italic=True)
    intro_p.paragraph_format.space_after = Pt(12)
    
    # Helper to add section heading
    def add_section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(clean_text(text))
        apply_text_formatting(run, size_pt=12, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        
    # Helper to add bullet points or paragraphs
    def add_body_paragraph(label, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run_label = p.add_run(f"{clean_text(label)}: ")
        apply_text_formatting(run_label, bold=True)
        run_text = p.add_run(clean_text(text))
        apply_text_formatting(run_text)
        
    # 1. Linguistic Fingerprint
    add_section_heading("1. Linguistic Fingerprint (Idiolect)")
    lf = data.get("linguistic_fingerprint", {})
    add_body_paragraph("Vocabulary & Syntax", lf.get("vocabulary_syntax", "N/A"))
    add_body_paragraph("Punctuation & Formatting", lf.get("punctuation_formatting", "N/A"))
    add_body_paragraph("Conversational Pacing", lf.get("conversational_pacing", "N/A"))
    
    # 2. Cognitive & Behavioral Patterns
    add_section_heading("2. Cognitive & Behavioral Patterns")
    cp = data.get("cognitive_patterns", {})
    add_body_paragraph("Decision-Making Style", cp.get("decision_making", "N/A"))
    add_body_paragraph("Cognitive Biases", cp.get("cognitive_biases", "N/A"))
    add_body_paragraph("Stress Response", cp.get("stress_response", "N/A"))
    add_body_paragraph("Intellectual Depth", cp.get("intellectual_depth", "N/A"))
    
    # 3. Value Systems & Worldview
    add_section_heading("3. Value Systems & Worldview")
    vs = data.get("value_systems", {})
    add_body_paragraph("Core Motivations", vs.get("core_motivations", "N/A"))
    add_body_paragraph("Core Fears", vs.get("core_fears", "N/A"))
    add_body_paragraph("Beliefs & Worldview", vs.get("beliefs_worldview", "N/A"))
    add_body_paragraph("Social Role", vs.get("social_role", "N/A"))
    
    # 4. Interests & Expertise
    add_section_heading("4. Interests & Expertise")
    ie = data.get("interests_expertise", {})
    add_body_paragraph("Domains of Knowledge", ie.get("domains", "N/A"))
    add_body_paragraph("Aesthetic Preferences", ie.get("aesthetic", "N/A"))
    
    # 5. Social Dynamics
    add_section_heading("5. Social Dynamics")
    sd = data.get("social_dynamics", {})
    add_body_paragraph("Interaction Style", sd.get("interaction_style", "N/A"))
    add_body_paragraph("Relationship Dynamics", sd.get("relationship_dynamics", "N/A"))
    
    # 6. Persona Engine (LLM System Prompt)
    add_section_heading("6. Persona Engine (LLM System Prompt)")
    prompt_p = doc.add_paragraph()
    prompt_run = prompt_p.add_run("Use the following system prompt to prime any LLM to act as this character with high fidelity:")
    apply_text_formatting(prompt_run, italic=True)
    
    # Create a callout box (table with 1 cell) for the system prompt
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=144, bottom=144, left=144, right=144) # ~10pt padding
    
    # Add light grey background to the cell
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), 'F2F2F2')
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Add border to the cell (left border only, thick grey)
    tcBorders = OxmlElement('w:tcBorders')
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '24') # 3pt
    left_border.set(qn('w:space'), '0')
    left_border.set(qn('w:color'), '7F7F7F')
    tcBorders.append(left_border)
    # Clear other borders
    for b in ['top', 'bottom', 'right']:
        border = OxmlElement(f'w:{b}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    cell._tc.get_or_add_tcPr().append(tcBorders)
    
    # Populate the cell with the system prompt
    prompt_text = generate_system_prompt_text(data)
    cell_p = cell.paragraphs[0]
    cell_p.paragraph_format.line_spacing = 1.15
    cell_p.paragraph_format.space_after = Pt(0)
    cell_run = cell_p.add_run(clean_text(prompt_text))
    apply_text_formatting(cell_run, font_name="Consolas", size_pt=9.5)
    
    doc.save(output_path)

def generate_system_prompt_text(data):
    name = data.get("character_name", "Unnamed Character")
    lf = data.get("linguistic_fingerprint", {})
    cp = data.get("cognitive_patterns", {})
    vs = data.get("value_systems", {})
    ie = data.get("interests_expertise", {})
    sd = data.get("social_dynamics", {})
    
    prompt = f"""You are acting as the fictional character: {name}.
Your goal is to simulate this character with absolute fidelity, matching their linguistic, cognitive, and behavioral patterns perfectly.

### 1. LINGUISTIC FINGERPRINT (HOW YOU TALK)
- Vocabulary & Syntax: {lf.get('vocabulary_syntax', 'N/A')}
- Punctuation & Formatting: {lf.get('punctuation_formatting', 'N/A')}
- Conversational Pacing: {lf.get('conversational_pacing', 'N/A')}

### 2. COGNITIVE & BEHAVIORAL PATTERNS (HOW YOU THINK)
- Decision-Making Style: {cp.get('decision_making', 'N/A')}
- Cognitive Biases: {cp.get('cognitive_biases', 'N/A')}
- Stress Response: {cp.get('stress_response', 'N/A')}
- Intellectual Depth: {cp.get('intellectual_depth', 'N/A')}

### 3. VALUE SYSTEMS & WORLDVIEW (WHAT DRIVES YOU)
- Core Motivations: {vs.get('core_motivations', 'N/A')}
- Core Fears: {vs.get('core_fears', 'N/A')}
- Beliefs & Worldview: {vs.get('beliefs_worldview', 'N/A')}
- Social Role: {vs.get('social_role', 'N/A')}

### 4. INTERESTS & EXPERTISE (WHAT YOU KNOW)
- Domains of Knowledge: {ie.get('domains', 'N/A')}
- Aesthetic Preferences: {ie.get('aesthetic', 'N/A')}

### 5. SOCIAL DYNAMICS (HOW YOU INTERACT)
- Interaction Style: {sd.get('interaction_style', 'N/A')}
- Relationship Dynamics: {sd.get('relationship_dynamics', 'N/A')}

### INSTRUCTIONS FOR SIMULATION:
1. Never break character. Do not refer to yourself as an AI or assistant.
2. Adopt the exact linguistic fingerprint, including code-switching, punctuation quirks, and pacing.
3. Respond to situations and questions using the character's cognitive patterns and value systems.
4. If asked about topics outside your expertise, respond as the character would (either with curiosity, dismissal, or redirection).
"""
    return prompt

def main():
    parser = argparse.ArgumentParser(description="Generate Persona Profile and LLM System Prompt.")
    parser.add_argument("--input", required=True, help="Path to input JSON file.")
    parser.add_argument("--output-dir", required=True, help="Directory to save output files.")
    args = parser.parse_args()
    
    if not os.path.exists(args.output_dir):
        os.makedirs(args.output_dir)
        
    with open(args.input, 'r') as f:
        data = json.load(f)
        
    char_name_clean = data.get("character_name", "character").lower().replace(" ", "_")
    
    # Generate DOCX
    docx_path = os.path.join(args.output_dir, f"{char_name_clean}_profile.docx")
    create_docx_profile(data, docx_path)
    print(f"Successfully generated DOCX profile: {docx_path}")
    
    # Generate Markdown Profile
    md_path = os.path.join(args.output_dir, f"{char_name_clean}_profile.md")
    with open(md_path, 'w') as f:
        f.write(f"# Deep Behavioral Profile: {data.get('character_name')}\n\n")
        f.write("## 1. Linguistic Fingerprint\n")
        for k, v in data.get("linguistic_fingerprint", {}).items():
            f.write(f"- **{k.replace('_', ' ').title()}**: {v}\n")
        f.write("\n## 2. Cognitive & Behavioral Patterns\n")
        for k, v in data.get("cognitive_patterns", {}).items():
            f.write(f"- **{k.replace('_', ' ').title()}**: {v}\n")
        f.write("\n## 3. Value Systems & Worldview\n")
        for k, v in data.get("value_systems", {}).items():
            f.write(f"- **{k.replace('_', ' ').title()}**: {v}\n")
        f.write("\n## 4. Interests & Expertise\n")
        for k, v in data.get("interests_expertise", {}).items():
            f.write(f"- **{k.replace('_', ' ').title()}**: {v}\n")
        f.write("\n## 5. Social Dynamics\n")
        for k, v in data.get("social_dynamics", {}).items():
            f.write(f"- **{k.replace('_', ' ').title()}**: {v}\n")
            
    print(f"Successfully generated Markdown profile: {md_path}")
    
    # Generate System Prompt Text File
    prompt_path = os.path.join(args.output_dir, f"{char_name_clean}_system_prompt.txt")
    with open(prompt_path, 'w') as f:
        f.write(generate_system_prompt_text(data))
    print(f"Successfully generated System Prompt: {prompt_path}")

if __name__ == "__main__":
    main()
